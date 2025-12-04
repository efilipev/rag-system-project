"""
Document Upload Service

Handles document upload, parsing, chunking, embedding, and storage in Qdrant.
Supports PDF, TXT, and DOCX files.
"""

import hashlib
import io
import re
import uuid
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Dict, List, Optional

from qdrant_client.http import models

from src.core.config import settings
from src.core.logging import logger
from src.services.document_processor import (
    DocumentProcessor,
    DocumentType,
    ProcessedDocument,
    get_document_processor,
)


@dataclass
class UploadResult:
    """Result of document upload operation"""
    success: bool
    document_id: str
    collection_name: str
    filename: str
    chunks_created: int
    processing_time_ms: float
    error: Optional[str] = None


@dataclass
class PDFParseResult:
    """Result of PDF parsing with page tracking"""
    text: str
    total_pages: int
    page_boundaries: List[tuple]  # List of (start_char, end_char, page_num)


class FileParser:
    """Parse different file types to extract text"""

    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """Extract text from PDF file (returns text only for backward compatibility)"""
        result = FileParser.parse_pdf_with_pages(file_content)
        return result.text

    @staticmethod
    def parse_pdf_with_pages(file_content: bytes) -> PDFParseResult:
        """Extract text from PDF file with page boundary tracking"""
        try:
            from pypdf import PdfReader

            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)

            text_parts = []
            page_boundaries = []
            current_pos = 0

            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                if text:
                    # Aggressive cleanup for PDFs with word-per-line extraction
                    cleaned = ' '.join(text.split())

                    # Try to restore paragraph breaks at sentence endings followed by capitals
                    cleaned = re.sub(r'([.!?])\s+([A-Z])', r'\1\n\n\2', cleaned)

                    # Also break on common section markers
                    cleaned = re.sub(r'\s*(Chapter \d+|Section \d+|\d+\.\d+)', r'\n\n\1', cleaned)

                    # Track page boundary
                    start_pos = current_pos
                    end_pos = current_pos + len(cleaned)
                    page_boundaries.append((start_pos, end_pos, page_num))

                    text_parts.append(cleaned)
                    current_pos = end_pos + 2  # +2 for "\n\n" separator

            return PDFParseResult(
                text="\n\n".join(text_parts),
                total_pages=len(reader.pages),
                page_boundaries=page_boundaries
            )
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise ValueError(f"Failed to parse PDF: {e}")

    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document

            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")

    @staticmethod
    def parse_txt(file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Unable to decode text file")
        except Exception as e:
            logger.error(f"Error parsing TXT: {e}")
            raise ValueError(f"Failed to parse text file: {e}")

    @staticmethod
    def parse_markdown(file_content: bytes) -> str:
        """Extract text from Markdown file"""
        return FileParser.parse_txt(file_content)

    @classmethod
    def parse(cls, file_content: bytes, content_type: str, filename: str) -> tuple[str, DocumentType, Optional[PDFParseResult]]:
        """
        Parse file content based on content type or filename extension

        Returns:
            Tuple of (extracted_text, document_type, pdf_result or None)
        """
        # Determine file type from content_type or extension
        extension = filename.lower().split('.')[-1] if '.' in filename else ''

        if content_type == 'application/pdf' or extension == 'pdf':
            pdf_result = cls.parse_pdf_with_pages(file_content)
            return pdf_result.text, DocumentType.PDF, pdf_result
        elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'] or extension == 'docx':
            return cls.parse_docx(file_content), DocumentType.TEXT, None
        elif content_type == 'text/markdown' or extension in ['md', 'markdown']:
            return cls.parse_markdown(file_content), DocumentType.MARKDOWN, None
        elif content_type.startswith('text/') or extension in ['txt', 'text']:
            return cls.parse_txt(file_content), DocumentType.TEXT, None
        else:
            # Try to parse as text
            try:
                return cls.parse_txt(file_content), DocumentType.TEXT, None
            except:
                raise ValueError(f"Unsupported file type: {content_type} ({extension})")


class UploadService:
    """
    Service for handling document uploads

    Flow:
    1. Parse uploaded file (PDF, DOCX, TXT)
    2. Chunk the document
    3. Generate embeddings
    4. Create/update Qdrant collection
    5. Store vectors with metadata
    """

    def __init__(self, vector_store):
        self.vector_store = vector_store
        self.document_processor = get_document_processor()
        self.file_parser = FileParser()

    def _sanitize_collection_name(self, name: str) -> str:
        """Sanitize collection name for Qdrant with _doc suffix"""
        # Remove file extension
        name = re.sub(r'\.[^.]+$', '', name)
        # Replace invalid characters with underscores
        name = re.sub(r'[^a-zA-Z0-9_-]', '_', name)
        # Ensure it starts with a letter
        if name and not name[0].isalpha():
            name = 'doc_' + name
        # Limit length (accounting for _doc suffix)
        name = name[:46]
        # Add _doc suffix for document collections
        return name.lower() + "_doc"

    async def _ensure_collection_exists(self, collection_name: str) -> None:
        """Create Qdrant collection if it doesn't exist"""
        try:
            collections = self.vector_store.qdrant_client.get_collections().collections
            collection_names = [col.name for col in collections]

            if collection_name not in collection_names:
                logger.info(f"Creating new collection: {collection_name}")
                self.vector_store.qdrant_client.create_collection(
                    collection_name=collection_name,
                    vectors_config=models.VectorParams(
                        size=settings.VECTOR_DIMENSION,
                        distance=models.Distance.COSINE,
                    ),
                )
                logger.info(f"Collection '{collection_name}' created successfully")
        except Exception as e:
            logger.error(f"Error ensuring collection exists: {e}")
            raise

    async def upload_document(
        self,
        file_content: bytes,
        filename: str,
        content_type: str,
        collection_name: Optional[str] = None,
        session_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> UploadResult:
        """
        Upload and process a document

        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME type
            collection_name: Target collection (auto-generated if not provided)
            session_id: Session ID for tracking
            metadata: Additional metadata

        Returns:
            UploadResult with processing details
        """
        start_time = datetime.now()
        document_id = str(uuid.uuid4())

        try:
            # Step 1: Parse file
            logger.info(f"Parsing file: {filename} ({content_type})")
            text, doc_type, pdf_result = self.file_parser.parse(file_content, content_type, filename)

            if not text.strip():
                return UploadResult(
                    success=False,
                    document_id=document_id,
                    collection_name="",
                    filename=filename,
                    chunks_created=0,
                    processing_time_ms=0,
                    error="No text content extracted from file"
                )

            logger.info(f"Extracted {len(text)} characters from {filename}")

            # Step 2: Determine collection name
            if not collection_name:
                collection_name = self._sanitize_collection_name(filename)

            # Step 3: Process document (chunk)
            logger.info(f"Processing document into chunks...")
            base_metadata = {
                "filename": filename,
                "content_type": content_type,
                "session_id": session_id,
                "upload_time": datetime.now().isoformat(),
                **(metadata or {})
            }

            processed_doc: ProcessedDocument = self.document_processor.process(
                text=text,
                title=filename,
                source=filename,
                document_type=doc_type,
                metadata=base_metadata
            )

            logger.info(f"Created {processed_doc.total_chunks} chunks")

            # Step 4: Ensure collection exists
            await self._ensure_collection_exists(collection_name)

            # Step 5: Generate embeddings and store
            logger.info(f"Generating embeddings and storing in collection: {collection_name}")

            # Helper function to find page number for a chunk based on its position
            def get_page_for_chunk(chunk_start: int, chunk_end: int) -> Optional[int]:
                if pdf_result is None:
                    return None
                # Find which page this chunk belongs to (based on majority overlap)
                for page_start, page_end, page_num in pdf_result.page_boundaries:
                    # Check if chunk overlaps with this page
                    overlap_start = max(chunk_start, page_start)
                    overlap_end = min(chunk_end, page_end)
                    if overlap_start < overlap_end:
                        return page_num
                # Fallback: estimate based on position
                if pdf_result.total_pages > 0 and len(text) > 0:
                    return int((chunk_start / len(text)) * pdf_result.total_pages) + 1
                return None

            texts = [chunk.text for chunk in processed_doc.chunks]
            metadatas = []
            for chunk in processed_doc.chunks:
                chunk_metadata = {
                    "document_id": document_id,
                    "chunk_id": chunk.chunk_id,
                    "chunk_index": chunk.chunk_index,
                    "total_chunks": chunk.total_chunks,
                    "title": filename,
                    "source": filename,
                    "has_math": chunk.has_math,
                    "has_code": chunk.has_code,
                    "word_count": chunk.word_count,
                    **base_metadata
                }
                # Add page number for PDFs
                if pdf_result is not None:
                    page_num = get_page_for_chunk(chunk.start_char, chunk.end_char)
                    if page_num:
                        chunk_metadata["page"] = page_num
                        chunk_metadata["total_pages"] = pdf_result.total_pages
                metadatas.append(chunk_metadata)

            # Generate embeddings
            embeddings = self.vector_store.embeddings.embed_documents(texts)

            # Create points for Qdrant
            points = [
                models.PointStruct(
                    id=str(uuid.uuid4()),
                    vector=embedding,
                    payload={
                        "text": text,
                        **metadata
                    }
                )
                for text, embedding, metadata in zip(texts, embeddings, metadatas)
            ]

            # Upsert to Qdrant
            self.vector_store.qdrant_client.upsert(
                collection_name=collection_name,
                points=points
            )

            processing_time = (datetime.now() - start_time).total_seconds() * 1000

            logger.info(f"Successfully uploaded {filename} to collection '{collection_name}' "
                       f"({len(points)} vectors, {processing_time:.0f}ms)")

            return UploadResult(
                success=True,
                document_id=document_id,
                collection_name=collection_name,
                filename=filename,
                chunks_created=len(points),
                processing_time_ms=processing_time
            )

        except Exception as e:
            processing_time = (datetime.now() - start_time).total_seconds() * 1000
            logger.error(f"Error uploading document: {e}", exc_info=True)
            return UploadResult(
                success=False,
                document_id=document_id,
                collection_name=collection_name or "",
                filename=filename,
                chunks_created=0,
                processing_time_ms=processing_time,
                error=str(e)
            )

    async def delete_document(
        self,
        document_id: str,
        collection_name: str
    ) -> bool:
        """
        Delete a document from a collection

        Args:
            document_id: Document ID to delete
            collection_name: Collection containing the document

        Returns:
            True if successful
        """
        try:
            # Delete all points with this document_id
            self.vector_store.qdrant_client.delete(
                collection_name=collection_name,
                points_selector=models.FilterSelector(
                    filter=models.Filter(
                        must=[
                            models.FieldCondition(
                                key="document_id",
                                match=models.MatchValue(value=document_id)
                            )
                        ]
                    )
                )
            )
            logger.info(f"Deleted document {document_id} from collection {collection_name}")
            return True
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            return False

    async def delete_collection(self, collection_name: str) -> bool:
        """
        Delete an entire collection

        Args:
            collection_name: Collection to delete

        Returns:
            True if successful
        """
        try:
            self.vector_store.qdrant_client.delete_collection(collection_name)
            logger.info(f"Deleted collection: {collection_name}")
            return True
        except Exception as e:
            logger.error(f"Error deleting collection: {e}")
            return False
